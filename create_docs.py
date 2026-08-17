import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = docx.Document()

# Set standard margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Color Palette Constants (Emerald Theme)
PRIMARY_COLOR = RGBColor(6, 78, 59)     # Emerald 900
SECONDARY_COLOR = RGBColor(16, 185, 129) # Emerald 500
DARK_TEXT = RGBColor(15, 23, 42)        # Slate 900
MUTED_TEXT = RGBColor(51, 65, 85)       # Slate 700

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Title Header
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = title_p.add_run('awaaz.ai')
run_title.font.size = Pt(28)
run_title.font.bold = True
run_title.font.color.rgb = PRIMARY_COLOR

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = sub_p.add_run('Every Voice Heard. Every Issue Resolved.\nTechnical, Operational & System Architecture Documentation')
run_sub.font.size = Pt(13)
run_sub.font.bold = True
run_sub.font.color.rgb = SECONDARY_COLOR

doc.add_paragraph()

# 1. Executive Summary
h1 = doc.add_heading(level=1)
r_h1 = h1.add_run('1. Executive Overview & System Architecture')
r_h1.font.color.rgb = PRIMARY_COLOR

p1 = doc.add_paragraph('awaaz.ai is an advanced, enterprise-grade AI-powered municipal grievance triage and predictive city operations platform. It bridges the gap between urban citizens and municipal administration through multi-lingual voice intake (English, Hindi, Marathi), Explainable AI (XAI) multi-class classification, 60-second agentic work order dispatch, real-time Google Maps spatial telemetry, 3-citizen verification protocols, and SHA-256 cryptographic audit ledgers.')
p1.style.font.size = Pt(10.5)

# 2. Comprehensive Technology Mapping Table
h2 = doc.add_heading(level=1)
r_h2 = h2.add_run('2. Technology Stack & Operational Mapping Table')
r_h2.font.color.rgb = PRIMARY_COLOR

p_table_intro = doc.add_paragraph('The table below provides a granular breakdown of every technology component, its codebase location, operational execution model, and security mechanism:')
p_table_intro.style.font.size = Pt(10.5)

table_data = [
    ['Layer / Component', 'Technology Used', 'Codebase Location', 'Operational Function & Execution Model', 'Security / Verification'],
    ['User Interface (Frontend)', 'React 18, Vite 5, Tailwind CSS v3', 'client/src/', 'Single-Page Application (SPA) providing responsive light-emerald dashboard UI, dynamic forms, and real-time state persistence.', 'Role-Based Component Guards (Citizen vs Officer)'],
    ['Brand & Visual Assets', 'Custom Transparent PNG Logo Emblem', 'client/public/logo.png, client/src/assets/logo.png', 'Transparent soundwave line + 3D letter A logo integrated in Navbar, Login Gateway, and Footer.', 'HTTPS CDN delivery'],
    ['Routing & Role Security', 'React Router v6, AuthContext', 'client/src/App.jsx, client/src/context/AuthContext.jsx', 'Client-side route guards (ProtectedRoute, OfficerRoute, CitizenRoute) enforcing role isolation.', 'LocalStorage session token persistence (civic_user)'],
    ['Voice & Speech AI Intake', 'Web Speech API (STT / TTS)', 'client/src/components/VoiceInput.jsx', 'Native speech recognition in EN, HI, MR. Auto-converts voice complaints into structured text with intent extraction.', 'Client-side local audio processing'],
    ['Privacy Shield & Object Blur', 'YOLOv8 Computer Vision + Regex Engine', 'client/src/utils/imageAnonymizer.js, client/src/components/ImageUpload.jsx', 'Auto-detects and blurs human faces & vehicle license plates on evidence photos. Redacts Aadhaar & phone numbers.', 'DPDP Act 2023 Compliant • Automatic PII Anonymization'],
    ['Location & Geocoding', 'Google Maps API, Geolocation', 'client/src/components/LocationPicker.jsx', 'Pins exact latitude/longitude coordinates and resolves addresses to Nagpur Municipal City Zones.', 'API Key restricted domain referrer'],
    ['Explainable AI (XAI) Triage', 'Multi-Class Priority Engine', 'client/src/components/XAIPanel.jsx, server/controllers/complaintController.js', 'Evaluates text keywords, proximity to schools/hospitals, and historical repair precedents (91-96% confidence).', 'Rule-based validation & Human Override controls'],
    ['Officer Work Order Copilot', 'Agentic Dispatch Engine', 'client/src/components/ResolutionCopilot.jsx', 'Autonomous 60-second contractor lookup, cost estimation, and work order generation.', 'Human-in-the-loop officer approval'],
    ['Work Resolution & Proof', 'Admin Photo Upload & Canvas Modal', 'client/src/components/ResolutionProofModal.jsx', 'Pops up when officer marks ticket solved. Uploads repair photo proof with GPS timestamp and engineering notes.', 'Required before status transition to Pending Verification'],
    ['Citizen Audit Verification', '3-Citizen Audit Engine', 'client/src/components/CitizenVerificationPanel.jsx, server/controllers/complaintController.js', 'Requires 3 independent citizens to inspect photo proof and verify work authenticity before final resolution.', '7-Day Verification Lock Window (reverts if unverified)'],
    ['City Digital Twin Map', '@react-google-maps/api, Dynamic Telemetry', 'client/src/components/DigitalTwinMap.jsx, client/src/pages/DigitalTwinPage.jsx', 'Real-time spatial simulation of municipal infrastructure health across Laxmi Nagar, Dharampeth, Sadar & Sitabuldi.', 'IoT sensor simulation & live status pins'],
    ['Infrastructure Heatmap', 'Chart.js, Custom Metric Gauges', 'client/src/components/HeatMap.jsx, client/src/pages/AnalyticsPage.jsx', 'Visualizes city zone failure hotspots for roads, water mainlines, sanitation dumps, and smart lighting.', 'City-wide operations analytics'],
    ['Backend API Server', 'Node.js, Express.js (REST)', 'server/index.js, server/routes/complaints.js', 'Handles HTTP requests for complaint management, user authentication, status updates, and verification endpoints.', 'CORS protection, Helmet security headers'],
    ['Database & Data Persistence', 'MongoDB (Mongoose) + JSON Fallback', 'server/models/Complaint.js, data/sample_complaints.json', 'Dual-mode database architecture. Uses MongoDB if connected, otherwise seamlessly falls back to persistent JSON storage.', 'Mongoose Schema validation'],
    ['Authentication & Security', 'JWT (JSON Web Tokens), bcryptjs', 'server/controllers/authController.js, server/models/User.js', 'Encrypted password hashing (salt 10) and signed JWT tokens with 24-hour expiration.', 'OFFICER_SECRET_KEY required for Officer/Admin signup'],
    ['Environment Configuration', 'dotenv, .env config files', 'server/.env, client/.env', 'Stores sensitive server secrets (OFFICER_SECRET_KEY, JWT_SECRET, MONGODB_URI).', 'Git-ignored (.env, server/.env, client/.env)'],
    ['Blockchain Audit Trail', 'SHA-256 Cryptographic Hash Engine', 'server/services/blockchainService.js, client/src/components/BlockchainAudit.jsx', 'Generates tamper-evident SHA-256 cryptographic block hashes for every grievance submission and status update.', 'Immutable ledger logging']
]

table = doc.add_table(rows=len(table_data), cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

col_widths = [Inches(1.2), Inches(1.3), Inches(1.5), Inches(2.2), Inches(1.4)]

for i, row in enumerate(table.rows):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
    if i == 0:
        trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

    for j, cell in enumerate(row.cells):
        cell.width = col_widths[j]
        cell.paragraphs[0].text = table_data[i][j]
        p = cell.paragraphs[0]
        p.style.font.size = Pt(8.5)

        if i == 0:
            set_cell_background(cell, '064E3B')
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
        else:
            if i % 2 == 1:
                set_cell_background(cell, 'F0FDF4')
            else:
                set_cell_background(cell, 'FFFFFF')
            for r in p.runs:
                r.font.color.rgb = DARK_TEXT

doc.add_paragraph()

# 3. Operational Workflows
h3 = doc.add_heading(level=1)
r_h3 = h3.add_run('3. End-to-End Operational Lifecycle & Workflows')
r_h3.font.color.rgb = PRIMARY_COLOR

workflows = [
    ('Stage 1: Resident Complaint Intake', 'Citizens submit grievances via voice (EN/HI/MR) or text. Location is pinned using GPS or manually selected from Nagpur Municipal City Zones. PII (phone/Aadhaar) is automatically redacted via Privacy Shield.'),
    ('Stage 2: XAI Triage & Priority Scoring', 'Explainable AI analyzes keywords, spatial proximity to safety zones (schools/hospitals), and historical repair trends to assign confidence scores (91-96%) and SLA resolution windows (e.g. 48h).'),
    ('Stage 3: Officer Triage & Agentic Dispatch', 'Authorized Officers view tickets in an interactive 5-stage Kanban flow (New -> Assigned -> In Progress -> Pending Verification -> Resolved). The Copilot generates automated contractor work orders.'),
    ('Stage 4: Work Completion & Photo Proof Submission', 'When an officer marks a ticket solved, the Resolution Proof Modal prompts for repair photo proof (GPS tagged) and engineering completion notes.'),
    ('Stage 5: 3-Citizen Verification & 7-Day Window', 'The ticket transitions to "Pending Verification" and triggers a 7-Day Verification Lock. 3 independent citizens must inspect and audit the photo proof. Once 3 verifications are recorded, it advances to "Verified & Resolved".'),
    ('Stage 6: Digital Twin Sync & Cryptographic Logging', 'All state changes update the live Google Maps City Digital Twin spatial telemetry map and append an immutable SHA-256 block hash to the public blockchain audit trail.')
]

for stage_title, stage_desc in workflows:
    p_st = doc.add_paragraph()
    r_st = p_st.add_run(stage_title)
    r_st.font.bold = True
    r_st.font.size = Pt(11)
    r_st.font.color.rgb = PRIMARY_COLOR
    
    p_sd = doc.add_paragraph(stage_desc)
    p_sd.style.font.size = Pt(10)

# 4. API Endpoints Table
doc.add_paragraph()
h4 = doc.add_heading(level=1)
r_h4 = h4.add_run('4. API REST Specification Table')
r_h4.font.color.rgb = PRIMARY_COLOR

api_data = [
    ['HTTP Method', 'API Endpoint', 'Payload / Query', 'Functionality & Security'],
    ['GET', '/api/complaints', 'None', 'Retrieves all grievances sorted by timestamp. Used for Kanban and Digital Twin telemetry.'],
    ['POST', '/api/complaints', 'title, category, location, description', 'Registers a new grievance. Applies PII redaction and computes SHA-256 audit hash.'],
    ['GET', '/api/complaints/:id', 'None', 'Fetches detailed complaint record including XAI rationale and verification history.'],
    ['PATCH', '/api/complaints/:id/status', 'status, resolutionProof, notes', 'Updates status. Triggers photo proof requirement for Pending Verification.'],
    ['POST', '/api/complaints/:id/verify', 'citizenName, comment', 'Records citizen authenticity audit (1/3, 2/3, 3/3 -> Verified & Resolved).'],
    ['POST', '/api/auth/register', 'name, mobile, email, secretKey, role', 'Registers user. Requires OFFICER_SECRET_KEY for Officer/Admin role authorization.'],
    ['POST', '/api/auth/login', 'identifier, password', 'Authenticates user and returns 24-hour signed JWT token.']
]

table_api = doc.add_table(rows=len(api_data), cols=4)
table_api.alignment = WD_TABLE_ALIGNMENT.CENTER

col_widths_api = [Inches(1.2), Inches(2.0), Inches(2.0), Inches(2.3)]

for i, row in enumerate(table_api.rows):
    for j, cell in enumerate(row.cells):
        cell.width = col_widths_api[j]
        cell.paragraphs[0].text = api_data[i][j]
        p = cell.paragraphs[0]
        p.style.font.size = Pt(8.5)

        if i == 0:
            set_cell_background(cell, '064E3B')
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
        else:
            if i % 2 == 1:
                set_cell_background(cell, 'F0FDF4')
            else:
                set_cell_background(cell, 'FFFFFF')
            for r in p.runs:
                r.font.color.rgb = DARK_TEXT

docx_path = r'd:\yash\complaint_box\awaaz_ai_technical_operational_documentation.docx'
doc.save(docx_path)
print(f'SUCCESS: Document saved at {docx_path}')
